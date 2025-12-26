"""
Document Parser for Mock Exam Questions
Extracts questions from PDF and Word documents using AI
"""

import io
import json
import re
import logging
from typing import Any, Dict, List, Optional, Tuple

import httpx
import PyPDF2
import streamlit as st
from docx import Document

import config

# Set up logger
logger = logging.getLogger(__name__)

# OCR support flag
OCR_AVAILABLE = False
try:
    import easyocr
    import fitz  # PyMuPDF
    import numpy as np
    from PIL import Image

    OCR_AVAILABLE = True
except ImportError:
    pass


class DocumentParser:
    """Parse PDF/Word documents and extract mock exam questions using AI"""

    # Patterns that indicate missing context
    # NOTE: These patterns only apply when NO scenario is provided
    # If a scenario/term sheet exists, these patterns are allowed
    INCOMPLETE_PATTERNS = [
        r"^The basis of which",
        r"^This product",
        r"^According to the case study",
        r"^Based on the scenario",
        r"^In the example",
        r"^The product described",
        r"^Under this arrangement",
        r"^For this investment",
        r"^In this case",
        r"^The structure outlined",
        r"^The instrument mentioned",
        r"^The investment product",  # References specific product without context
        r"^The investment.*profile",  # References specific investment without context
        r"^Today is the maturity date",  # References specific maturity event without context
        r"the fund has not been auto-redeemed",  # References specific fund event
        r"basis.*mandatory\s*redemption.*determined",
        r"payout.*maturity.*determined by the\s*$",
        r"determined by the\s*Choices",
        r"^Referring to",
        r"^For the",
        r"^With reference to",
        r"^With the spot",  # "With the spot HSI trades..."
        r"^Assuming throughout",  # "Assuming throughout the entire period..."
        r"^Assuming that the",  # "Assuming that the HSI trades..."
        r"^With a (more|less|longer|shorter)",  # "With a longer investment tenure..."
    ]

    # OCR corruption patterns
    OCR_CORRUPTION_PATTERNS = [
        (r"\$[oO]\.", "Dollar sign followed by letter o/O instead of zero"),
        (r"\bordinai-y\b", "ordinai-y instead of ordinary"),
        (r"\bordinan/\b", "ordinan/ instead of ordinary"),
        (r"\bstmctures\b", "stmctures instead of structures"),
        (r"\binvestrnent\b", "investrnent instead of investment"),
        (r"\bgovernrnent\b", "governrnent instead of government"),
        (r"\bmanagernent\b", "managernent instead of management"),
        (r"\bperforrnance\b", "performrnance instead of performance"),
        (r"\brnake\b", "rnake instead of make"),
        (r"\brnay\b", "rnay instead of may"),
        (r"\brnust\b", "rnust instead of must"),
        (r"\brnore\b", "rnore instead of more"),
        (r"\b[Il]00\%", "Letter I or l instead of 1 in 100%"),
        (r"\s{4,}", "Four or more consecutive spaces"),
        (r"determined by the\s*Choices", "Question text bleeding into choices"),
    ]

    # Common spacing/typo patterns (UAT only - not production)
    # These catch typos that AI parsing or OCR missed
    TYPO_PATTERNS = [
        (r"\bajoint\b", "ajoint should be 'a joint'"),
        (r"\btobe\b", "tobe should be 'to be'"),
        (r"\bofthe\b", "ofthe should be 'of the'"),
        (r"\binthe\b", "inthe should be 'in the'"),
        (r"\bonthe\b", "onthe should be 'on the'"),
        (r"\batthe\b", "atthe should be 'at the'"),
        (r"\bforthe\b", "forthe should be 'for the'"),
        (r"\bwiththe\b", "withthe should be 'with the'"),
        (r"\bfromthe\b", "fromthe should be 'from the'"),
        (r"\bthat\s+the\s+the\b", "duplicate 'the the'"),
        (r"\ba\s+a\b", "duplicate 'a a'"),
        (r"\bwill\s+will\b", "duplicate 'will will'"),
        (r"\bis\s+is\b", "duplicate 'is is'"),
    ]

    # Multi-answer question patterns (questions requiring more than one correct answer)
    # These are NOT SUPPORTED and should be rejected during upload
    MULTI_ANSWER_PATTERNS = [
        r"\(select all that apply\)",
        r"\(select all\)",
        r"\(choose all that apply\)",
        r"\(choose all\)",
        r"\(tick all that apply\)",
        r"\(check all that apply\)",
        r"select all (the\s+)?options? that (are )?correct",
        r"select all (the\s+)?correct (options?|answers?|statements?)",
        r"choose all (the\s+)?options? that (are )?correct",
        r"which (of the following )?are correct",
        r"which (of the following )?statements? are (true|correct)",
        r"select (all|multiple) correct answer",
        r"mark all that apply",
        r"identify all (the\s+)?(correct|true)",
    ]

    def __init__(self):
        self.api_key = config.OPENROUTER_API_KEY
        self.model = config.OPENROUTER_MODEL
        self.base_url = config.OPENROUTER_BASE_URL

    def _detect_incomplete_question(self, question_text: str, scenario: Optional[str] = None) -> Optional[str]:
        """
        Detect if a question references missing context.
        Returns the reason if incomplete, None if complete.

        Args:
            question_text: The question text to validate
            scenario: Optional scenario/term sheet context. If provided, questions can reference "this", "the HSI", etc.
        """
        question_text = question_text.strip()

        # If a scenario is provided, the question can reference context (it's complete)
        # This allows questions like "What is the return of this product?" when a term sheet is attached
        if scenario and len(scenario.strip()) > 50:  # Meaningful scenario (not just a short note)
            return None

        # Check against patterns
        for pattern in self.INCOMPLETE_PATTERNS:
            if re.search(pattern, question_text, re.IGNORECASE):
                return "references missing context"

        # Check if question starts with article + noun without context
        if re.match(r"^The\s+\w+\s+(of|for|in)\s+which", question_text, re.IGNORECASE):
            return "incomplete sentence structure"

        # Check for references to "this/that product/fund/instrument" without context
        # Allow for adjectives like "this Structured fund", "that new product"
        if re.search(
            r"\b(this|that)\s+(?:\w+\s+)*(fund|product|instrument|security|option|structure)\b",
            question_text,
            re.IGNORECASE,
        ):
            return "references 'this/that' without context"

        return None

    def _detect_ocr_corruption(self, question_text: str) -> Optional[str]:
        """
        Detect OCR corruption in question text (spelling errors, malformed text).
        Returns the first corruption found, None if clean.
        """
        # Always check OCR corruption patterns
        for pattern, description in self.OCR_CORRUPTION_PATTERNS:
            if re.search(pattern, question_text, re.IGNORECASE):
                return description

        # Check typo patterns only in non-production environments
        if config.ENVIRONMENT != "production":
            for pattern, description in self.TYPO_PATTERNS:
                if re.search(pattern, question_text, re.IGNORECASE):
                    return description

        return None

    def _detect_multi_answer_question(self, question_text: str, correct_index: Any) -> Optional[str]:
        """
        Detect if a question requires multiple correct answers (not supported).
        Returns the reason if multi-answer detected, None if single-answer.

        Args:
            question_text: The question text to check
            correct_index: The correct_index field from AI (can be int or list)
        """
        # Check if correct_index is a list with multiple answers
        if isinstance(correct_index, list) and len(correct_index) > 1:
            return f"multiple correct answers provided ({len(correct_index)} answers)"

        # Check question text for multi-answer patterns
        for pattern in self.MULTI_ANSWER_PATTERNS:
            if re.search(pattern, question_text, re.IGNORECASE):
                return "contains 'select all that apply' or similar instruction"

        return None

    def parse_document(
        self, uploaded_file, pool_id=None, pool_name=None
    ) -> Tuple[bool, List[Dict[str, Any]], str]:
        """
        Parse uploaded document and extract questions

        Args:
            uploaded_file: The uploaded file object
            pool_id: Optional pool ID for background OCR processing
            pool_name: Optional pool name for background OCR processing

        Returns:
            (success, questions_list, error_message)
        """
        try:
            # Determine file type
            file_extension = uploaded_file.name.split(".")[-1].lower()

            # For Word documents, check if it contains JSON first
            if file_extension in ["docx", "doc"]:
                # Extract text to check for JSON
                text = self._extract_text_from_word(uploaded_file)
                text_stripped = text.strip()

                # Check if document contains pure JSON array
                if text_stripped.startswith('[') and text_stripped.endswith(']'):
                    st.info("📋 Detected JSON array in Word document, parsing directly...")
                    try:
                        json_data = json.loads(text_stripped)
                        if isinstance(json_data, list) and len(json_data) > 0:
                            # Convert to expected format
                            questions = []
                            for item in json_data:
                                if not isinstance(item, dict):
                                    continue

                                # Support both "choices"/"options" and "correct_index"/"correct_answer"
                                choices = item.get("choices") or item.get("options")
                                correct_index = item.get("correct_index")
                                if correct_index is None:
                                    correct_index = item.get("correct_answer")

                                if item.get("question") and choices and correct_index is not None:
                                    questions.append({
                                        "question": item["question"],
                                        "choices": choices,
                                        "correct_index": correct_index,
                                        "correct_answer": correct_index,
                                        "scenario": item.get("scenario"),
                                        "explanation_seed": item.get("explanation_seed")
                                    })

                            if questions:
                                st.success(f"✅ Parsed {len(questions)} questions directly from JSON")
                                valid_questions = self._validate_questions(questions)
                                if valid_questions:
                                    return True, valid_questions, ""
                    except json.JSONDecodeError as e:
                        st.warning(f"⚠️ JSON detection failed: {e}. Trying structured parsing...")

                # Not JSON or JSON parsing failed - try structured parsing
                st.info("🔍 Attempting structured parsing (faster and more accurate)...")
                uploaded_file.seek(0)  # Reset file pointer
                structured_questions = self._parse_word_structured(uploaded_file)

                if structured_questions:
                    st.success(
                        f"✅ Structured parsing successful! Extracted {len(structured_questions)} questions"
                    )
                    # Validate questions
                    valid_questions = self._validate_questions(structured_questions)
                    if valid_questions:
                        return True, valid_questions, ""
                    else:
                        st.warning(
                            "⚠️ Structured parsing validation failed, falling back to AI extraction..."
                        )

                else:
                    st.info("📄 Structured parsing found no questions, using AI extraction...")

            # Extract text from document for AI parsing
            if file_extension == "pdf":
                text = self._extract_text_from_pdf(uploaded_file, pool_id, pool_name)
            elif file_extension in ["docx", "doc"]:
                text = self._extract_text_from_word(uploaded_file)
            else:
                return False, [], f"Unsupported file type: {file_extension}"

            # Check if background processing was triggered
            if text == "__BACKGROUND_PROCESSING__":
                # Return success with empty list - questions will be added by background process
                return True, [], ""

            if not text or len(text.strip()) < 50:
                return False, [], "Could not extract sufficient text from document"

            # Note: For Word documents, JSON is already checked above. This fallback is for PDFs.
            # Use AI to parse questions from text
            st.info("🤖 Using AI to extract questions from document...")
            questions = self._parse_questions_with_ai(text)

            if not questions:
                return False, [], "No questions could be extracted from the document"

            # Validate questions
            valid_questions = self._validate_questions(questions)

            if not valid_questions:
                return False, [], "Extracted questions did not pass validation"

            return True, valid_questions, ""

        except Exception as e:
            return False, [], f"Error parsing document: {str(e)}"

    def _extract_text_from_pdf(self, uploaded_file, pool_id=None, pool_name=None) -> str:
        """Extract text from PDF file, with OCR fallback for scanned documents"""
        try:
            # Reset file pointer
            uploaded_file.seek(0)

            # Read PDF
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            num_pages = len(pdf_reader.pages)

            # Extract text from all pages
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

            # Check if we got meaningful text
            clean_text = text.strip()
            if len(clean_text) > 100:
                return clean_text

            # Text extraction failed or minimal text - this is a scanned PDF
            st.info(
                f"📄 PDF appears to be scanned ({num_pages} pages). Checking if background processing is needed..."
            )

            if not OCR_AVAILABLE:
                st.warning(
                    "⚠️ OCR libraries not available. Install easyocr and pymupdf for scanned PDF support."
                )
                return clean_text

            # If document is large (>25 pages), use background processing
            if num_pages > 25 and pool_id and pool_name:
                return self._trigger_background_ocr(uploaded_file, pool_id, pool_name, num_pages)

            # For smaller documents, process inline with OCR
            st.info(f"📄 Processing {num_pages} pages with OCR (this may take a few minutes)...")
            uploaded_file.seek(0)
            file_bytes = uploaded_file.read()
            ocr_text = self._ocr_pdf(file_bytes)

            if ocr_text and len(ocr_text.strip()) > 100:
                st.success(f"✅ OCR extracted {len(ocr_text)} characters from scanned PDF")
                return ocr_text
            else:
                st.warning("⚠️ OCR could not extract sufficient text")
                return clean_text

        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")

    def _trigger_background_ocr(self, uploaded_file, pool_id, pool_name, num_pages) -> str:
        """
        Save PDF and trigger background OCR processing for large documents

        Returns a special marker that indicates background processing was started
        """
        import os
        import subprocess
        import sys
        import tempfile

        try:
            # Save PDF to temporary file
            temp_dir = tempfile.gettempdir()
            temp_filename = f"ocr_pending_{pool_id}_{uploaded_file.name}"
            temp_path = os.path.join(temp_dir, temp_filename)

            uploaded_file.seek(0)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.read())

            # Trigger background OCR processor
            python_path = sys.executable
            script_path = os.path.join(os.getcwd(), "background_ocr_processor.py")

            # Log to file instead of suppressing output
            import time
            log_file_path = os.path.join(os.getcwd(), "background_ocr_processor.log")
            log_file = open(log_file_path, "a")

            try:
                process = subprocess.Popen(
                    [python_path, script_path, temp_path, pool_id, pool_name, uploaded_file.name],
                    stdout=log_file,
                    stderr=subprocess.STDOUT,  # Merge stderr into stdout
                    start_new_session=True,
                )

                # Verify process started successfully
                time.sleep(0.5)  # Give it a moment to start
                if process.poll() is not None:
                    # Process already died
                    raise Exception(f"Background OCR process failed to start (exit code: {process.returncode})")

            finally:
                # Always close log file handle
                log_file.close()

            # Create database status record for tracking
            try:
                from db import db
                # Use synchronous client to avoid event loop conflicts
                db.admin_client.table('background_upload_status').insert({
                    'pool_id': pool_id,
                    'pool_name': pool_name,
                    'status': 'processing',
                    'total_questions': 0,
                    'processed_questions': 0,
                    'error_message': None
                }).execute()
                st.info(f"📊 Background process tracking: Check 'Manage Pools' to monitor progress")
            except Exception as e:
                logger.warning(f"Failed to create background status record: {e}")

            st.success(
                f"✅ Large scanned PDF detected ({num_pages} pages)\n\n"
                f"📤 Background OCR processing started! Questions will be added to the pool automatically when processing completes.\n\n"
                f"⏱️ Estimated time: {num_pages * 15 // 60} minutes\n\n"
                f"💡 You can close this page and check back later - processing continues in the background."
            )

            # Return a special marker that tells the upload process to skip this file
            # (it will be processed in background)
            return "__BACKGROUND_PROCESSING__"

        except Exception as e:
            st.error(f"Failed to start background OCR processing: {e}")
            # Fall back to inline OCR with page limit
            uploaded_file.seek(0)
            file_bytes = uploaded_file.read()
            return self._ocr_pdf(file_bytes)

    def _ocr_pdf(self, file_bytes: bytes, max_pages: int = 25) -> str:
        """
        Extract text from scanned PDF using OCR

        Args:
            file_bytes: PDF file content
            max_pages: Maximum number of pages to process (to avoid timeouts)
        """
        try:
            # Open PDF with PyMuPDF
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            total_pages = len(pdf_doc)

            # Limit pages to avoid timeout
            pages_to_process = min(total_pages, max_pages)

            if total_pages > max_pages:
                st.warning(
                    f"⚠️ PDF has {total_pages} pages. Processing first {max_pages} pages to avoid timeout. "
                    f"For full document processing, please split into smaller files."
                )
            else:
                st.info(f"📄 Processing {total_pages} pages with OCR...")

            # Initialize EasyOCR reader (lazy load)
            reader = easyocr.Reader(["en"], gpu=False, verbose=False)

            all_text = []
            progress_bar = st.progress(0)

            for i in range(pages_to_process):
                page = pdf_doc[i]

                # Update progress
                progress_bar.progress(
                    (i + 1) / pages_to_process,
                    text=f"🔍 OCR processing page {i+1}/{pages_to_process}...",
                )

                # Render page to image (150 DPI - faster than 200, still good quality)
                mat = fitz.Matrix(150 / 72, 150 / 72)
                pix = page.get_pixmap(matrix=mat)

                # Convert to PIL Image then numpy array
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img_array = np.array(img)

                # Run OCR with optimized settings
                try:
                    results = reader.readtext(
                        img_array,
                        paragraph=True,  # Group text into paragraphs (faster)
                        batch_size=4,  # Process in small batches
                    )
                    # Extract text from results
                    page_text = "\n".join([result[1] for result in results])
                    all_text.append(page_text)
                except Exception as e:
                    st.warning(f"⚠️ OCR failed on page {i+1}: {str(e)}")
                    continue

            pdf_doc.close()
            progress_bar.empty()

            combined_text = "\n\n".join(all_text)

            if combined_text:
                st.success(
                    f"✅ OCR completed: {len(combined_text)} characters extracted from {pages_to_process} pages"
                )

            return combined_text

        except Exception as e:
            st.error(f"❌ OCR failed: {str(e)}")
            return ""

    def _extract_text_from_word(self, uploaded_file) -> str:
        """Extract text from Word document"""
        try:
            # Reset file pointer
            uploaded_file.seek(0)

            # Read Word document
            doc = Document(uploaded_file)
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading Word document: {str(e)}")

    def _parse_questions_with_ai(self, text: str) -> List[Dict[str, Any]]:
        """Use AI to parse questions from extracted text"""
        try:
            # If text is large, process in chunks
            chunk_size = 8000
            all_questions = []

            if len(text) <= chunk_size:
                # Small document - process all at once
                prompt = self._create_extraction_prompt(text)
                questions = self._call_ai_api(prompt)
                return questions
            else:
                # Large document - process in chunks
                st.info(f"📄 Document is large ({len(text)} chars). Processing in chunks...")

                # Split text into chunks (try to split at question boundaries)
                chunks = self._split_text_into_chunks(text, chunk_size)

                for i, chunk in enumerate(chunks, 1):
                    st.info(f"🔄 Processing chunk {i}/{len(chunks)}...")
                    prompt = self._create_extraction_prompt(chunk)
                    chunk_questions = self._call_ai_api(prompt)

                    if chunk_questions:
                        all_questions.extend(chunk_questions)
                        st.success(f"✅ Extracted {len(chunk_questions)} questions from chunk {i}")

                st.success(f"🎉 Total extracted: {len(all_questions)} questions from all chunks")
                return all_questions

        except Exception as e:
            st.error(f"AI parsing error: {str(e)}")
            return []

    def _split_text_into_chunks(self, text: str, chunk_size: int) -> List[str]:
        """Split text into chunks, trying to keep questions together"""
        chunks = []
        current_chunk = ""

        # Split by common question delimiters
        lines = text.split("\n")

        for line in lines:
            # Check if adding this line would exceed chunk size
            if len(current_chunk) + len(line) > chunk_size and current_chunk:
                # Save current chunk and start new one
                chunks.append(current_chunk)
                current_chunk = line + "\n"
            else:
                current_chunk += line + "\n"

        # Add last chunk
        if current_chunk.strip():
            chunks.append(current_chunk)

        return chunks

    def _create_extraction_prompt(self, text: str) -> str:
        """Create prompt for AI to extract questions"""
        return f"""You are a specialized question extraction system. Extract all multiple-choice questions from the following document text.

DOCUMENT TEXT:
{text}

INSTRUCTIONS:
1. Identify all multiple-choice questions in the text

2. **IMPORTANT - EXTRACT TERM SHEETS & CASE STUDIES**: Many exam documents contain:
   - Term sheets describing investment products (e.g., "HSI Range Accrual Note")
   - Case studies with product specifications
   - Scenario descriptions before a series of questions

   When you find such contextual information, extract it and attach it to ALL related questions in the "scenario" field.
   Questions that reference "this product", "the HSI", "this RAN", etc. are VALID if they have the term sheet in their scenario.

3. **CRITICAL - HANDLE LETTERED SUB-ITEMS IN QUESTIONS**:
   Many questions have this structure:
   ```
   31. Which of the following are elements of best execution?
   a. Speed of execution
   b. Price confidentiality
   c. Likelihood of execution
   d. Client instructions

   (a), (c) and (d)
   (a), (b) and (d)
   (b), (c) and (d)
   (a), (b), (c) and (d) [CORRECT]
   ```

   In this pattern:
   - The lettered items (a, b, c, d) are PART OF THE QUESTION TEXT, NOT answer choices
   - The actual answer choices are the combinations like "(a), (c) and (d)"

   You MUST include the lettered sub-items in the question text:
   ```
   "question": "Which of the following are elements of best execution?\n\na. Speed of execution\nb. Price confidentiality\nc. Likelihood of execution\nd. Client instructions"
   ```

4. **CRITICAL - FIX OCR ERRORS**: The text may contain OCR errors where words are split into individual letters with spaces or line breaks. Fix these errors:
   - Example: "p e r s h a r e" should be "per share"
   - Example: "r e c e i v e d" should be "received"
   - Example: "d i v i d e n d s o f" should be "dividends of"
   - Remove extra spaces between individual letters
   - Reconstruct proper words from letter sequences
   - Fix formatting issues like conjoined words (e.g., "pershareand" → "per share and")
   - Preserve mathematical symbols like $, %, numbers

5. For each question, extract:
   - The question text INCLUDING any lettered sub-items (a, b, c, d) that define what the answer choices refer to
   - All answer choices (which may reference the sub-items, e.g., "(a) and (b)")
   - The correct answer
   - **CRITICAL**: Any term sheet/case study/scenario that provides context for this question
   - Any explanation or additional notes

6. Format your response as a JSON array of questions. Each question should have:
   - "question": the question text INCLUDING lettered sub-items if present (string, OCR errors fixed)
   - "choices": array of answer choice texts (array of strings, OCR errors fixed)
   - "correct_index": zero-based index of correct answer (integer, 0-3)
   - "scenario": **IMPORTANT** - Include the full term sheet/case study/scenario text if the question references it (string or null)
   - "explanation_seed": optional hint for explanation (string or null)

EXAMPLE OUTPUT FORMAT (DO NOT COPY THESE EXAMPLES - EXTRACT ONLY FROM DOCUMENT):
[
  {{
    "question": "Which of the following are elements of best execution?\\n\\na. Speed of execution\\nb. Price confidentiality\\nc. Likelihood of execution\\nd. Client instructions",
    "choices": ["(a), (c) and (d)", "(a), (b) and (d)", "(b), (c) and (d)", "(a), (b), (c) and (d)"],
    "correct_index": 3,
    "scenario": null,
    "explanation_seed": null
  }},
  {{
    "question": "What is the maximum return of this product?",
    "choices": ["5%", "10%", "15%", "20%"],
    "correct_index": 2,
    "scenario": "ABC Range Accrual Note: Principal $100,000, Tenor: 2 years, Reference Index: HSI, Coupon: 7.5% p.a. if HSI stays within range, Maximum Return: 15% p.a.",
    "explanation_seed": "Term sheet specifies maximum return"
  }}
]

CRITICAL RULES:
- Extract ONLY questions that appear in the DOCUMENT TEXT above
- DO NOT include the example questions in your output
- DO NOT generate or hallucinate questions - extract only what exists in the document
- Return ONLY the JSON array, no other text
- correct_index must be 0-based (0, 1, 2, 3...)
- Include at least 2 choices per question
- If you cannot find clear questions in the document, return an empty array: []
- ONLY extract questions related to finance, investments, regulations, or business topics
- REJECT and do NOT extract general trivia, programming, geography, or unrelated topics

Extract all questions now:"""

    def _call_ai_api(self, prompt: str, max_retries: int = 2) -> List[Dict[str, Any]]:
        """Call OpenRouter API to extract questions with retry logic"""

        # Check if API key is configured (allow AI parsing even in DEMO_MODE)
        if not self.api_key or self.api_key == "demo" or self.api_key == "your_openrouter_api_key":
            # No valid API key configured
            st.warning(
                "⚠️ OpenRouter API key not configured. Using basic pattern matching fallback."
            )
            st.info(
                "💡 To enable AI-powered document parsing, add your OPENROUTER_API_KEY to .streamlit/secrets.toml"
            )
            return self._demo_parse_fallback()

        import time
        last_error = None

        for attempt in range(max_retries + 1):
            try:
                if attempt > 0:
                    wait_time = 2 ** attempt  # Exponential backoff: 2s, 4s
                    st.info(f"🔄 Retrying API call (attempt {attempt + 1}/{max_retries + 1}) after {wait_time}s...")
                    time.sleep(wait_time)

                headers = {
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://mockexamify.com",
                    "X-Title": "MockExamify",
                }

                payload = {
                    "model": self.model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,  # Low temperature for consistent extraction
                    "max_tokens": 4000,
                }

                # Make synchronous request
                with httpx.Client(timeout=60.0) as client:
                    response = client.post(
                        f"{self.base_url}/chat/completions", headers=headers, json=payload
                    )

                if response.status_code != 200:
                    raise Exception(f"API returned status {response.status_code}: {response.text}")

                result = response.json()

                # Check if response has expected structure
                if "choices" not in result:
                    logger.error(f"Unexpected API response structure: {result}")
                    raise Exception(f"API response missing 'choices' field. Response: {result.get('error', result)}")

                if not result["choices"] or len(result["choices"]) == 0:
                    raise Exception("API response has empty 'choices' array")

                if "message" not in result["choices"][0]:
                    raise Exception("API response missing 'message' field in choices[0]")

                if "content" not in result["choices"][0]["message"]:
                    raise Exception("API response missing 'content' field in message")

                content = result["choices"][0]["message"]["content"]

                # Extract JSON from response
                questions = self._extract_json_from_response(content)

                # Success - return questions
                return questions

            except Exception as e:
                last_error = e
                logger.warning(f"API call attempt {attempt + 1} failed: {str(e)}")

                # If this was the last attempt, fail
                if attempt == max_retries:
                    st.error(f"AI API call failed after {max_retries + 1} attempts: {str(last_error)}")
                    logger.error(f"AI API error details: {str(last_error)}", exc_info=True)
                    return []
                # Otherwise, continue to next retry

        # This should never be reached, but just in case
        return []

    def _extract_json_from_response(self, content: str) -> List[Dict[str, Any]]:
        """Extract JSON array from AI response"""
        try:
            # Try to parse directly
            questions = json.loads(content)
            return questions if isinstance(questions, list) else []
        except json.JSONDecodeError:
            # Try to find JSON array in the content
            json_match = re.search(r"\[[\s\S]*\]", content)
            if json_match:
                try:
                    questions = json.loads(json_match.group(0))
                    return questions if isinstance(questions, list) else []
                except json.JSONDecodeError:
                    pass
            return []

    def _demo_parse_fallback(self) -> List[Dict[str, Any]]:
        """Fallback parser for demo mode - extracts questions using regex patterns"""
        st.info("Using basic pattern matching. For better results, configure AI API key.")

        # Return empty list for now - user should provide API key for real parsing
        # This could be enhanced with regex-based parsing if needed
        return []

    def _auto_validate_calculations(
        self, question: str, choices: List[str], ai_index: int
    ) -> Tuple[int, bool]:
        """
        Auto-validate calculation questions and correct AI errors.
        Returns: (corrected_index, was_corrected)
        """
        import re

        question_lower = question.lower()

        # Only validate if all choices are numeric
        try:
            numeric_choices = [
                float(choice.replace(",", "").replace("$", "").strip()) for choice in choices
            ]
        except (ValueError, AttributeError):
            # Not all choices are numeric, skip validation
            return ai_index, False

        # Pattern 1: Cross Rate Calculation (GBP/USD × USD/CNY = GBP/CNY)
        cross_rate_match = re.search(
            r"(\w+)/(\w+)\s*[=:]\s*([\d.]+).*?(\w+)/(\w+)\s*[=:]\s*([\d.]+)", question
        )
        if cross_rate_match and "cross rate" in question_lower:
            try:
                # Extract: GBP/USD = 1.6824, USD/CNY = 6.2518
                rate1 = float(cross_rate_match.group(3))
                rate2 = float(cross_rate_match.group(6))
                calculated = rate1 * rate2

                # Find closest choice
                closest_index = min(
                    range(len(numeric_choices)), key=lambda i: abs(numeric_choices[i] - calculated)
                )

                if closest_index != ai_index:
                    return closest_index, True
            except (ValueError, AttributeError, KeyError):
                pass

        # Pattern 2: Sharpe Ratio
        sharpe_match = re.search(
            r"earned\s+([\d.]+)\s*%.*?standard deviation.*?([\d.]+)\s*%.*?risk[‑-]free.*?([\d.]+)\s*%",
            question,
        )
        if sharpe_match and "sharpe" in question_lower:
            try:
                portfolio_return = float(sharpe_match.group(1))
                std_dev = float(sharpe_match.group(2))
                risk_free = float(sharpe_match.group(3))

                sharpe_ratio = (portfolio_return - risk_free) / std_dev

                # Find closest choice
                closest_index = min(
                    range(len(numeric_choices)),
                    key=lambda i: abs(numeric_choices[i] - sharpe_ratio),
                )

                if closest_index != ai_index:
                    return closest_index, True
            except (ValueError, AttributeError, KeyError):
                pass

        # Pattern 3: Simple percentage calculation
        percentage_match = re.search(r"([\d.]+)\s*%.*?of.*?\$?([\d,]+)", question)
        if percentage_match and ("calculate" in question_lower or "what is" in question_lower):
            try:
                percentage = float(percentage_match.group(1)) / 100
                amount = float(percentage_match.group(2).replace(",", ""))
                calculated = percentage * amount

                # Find closest choice
                closest_index = min(
                    range(len(numeric_choices)), key=lambda i: abs(numeric_choices[i] - calculated)
                )

                if (
                    closest_index != ai_index
                    and abs(numeric_choices[closest_index] - calculated) < 0.01
                ):
                    return closest_index, True
            except (ValueError, AttributeError, KeyError):
                pass

        # No correction needed
        return ai_index, False

    def _is_finance_related(self, question_text: str, choices: List[str]) -> bool:
        """
        Check if a question is finance/investment/business related.
        Returns True if finance-related, False if general knowledge/trivia.
        """
        # Convert to lowercase for matching
        text_lower = (question_text + " " + " ".join(choices)).lower()

        # Finance/investment keywords (positive indicators)
        finance_keywords = [
            "investment",
            "portfolio",
            "fund",
            "dividend",
            "yield",
            "coupon",
            "bond",
            "equity",
            "stock",
            "share",
            "capital",
            "return",
            "risk",
            "maturity",
            "redemption",
            "price",
            "rate",
            "index",
            "basket",
            "derivative",
            "option",
            "warrant",
            "note",
            "structured",
            "security",
            "asset",
            "liability",
            "financial",
            "market",
            "trading",
            "broker",
            "investor",
            "client",
            "regulatory",
            "compliance",
            "mas",
            "cmfas",
            "sfa",
            "faa",
            "cacs",
            "interest",
            "principal",
            "credit",
            "default",
            "issuer",
            "counterparty",
            "underlying",
            "performance",
            "sharpe",
            "volatility",
            "correlation",
            "payout",
            "premium",
            "strike",
            "call",
            "put",
            "expiry",
            "settlement",
        ]

        # General knowledge keywords (negative indicators - red flags)
        general_knowledge_keywords = [
            "capital of",
            "programming language",
            "javascript",
            "python",
            "java",
            "france",
            "paris",
            "london",
            "geography",
            "history",
            "science",
            "chemistry",
            "physics",
            "biology",
            "mathematics",
            "algebra",
            "what is 2+2",
            "output of print",
            "capital city",
            "country",
            "president",
            "prime minister",
            "famous for",
            "invented by",
        ]

        # Check for general knowledge red flags
        for keyword in general_knowledge_keywords:
            if keyword in text_lower:
                return False

        # Check for finance keywords
        finance_match_count = sum(1 for keyword in finance_keywords if keyword in text_lower)

        # If we find multiple finance keywords, it's likely finance-related
        if finance_match_count >= 2:
            return True

        # If no strong signals either way, be conservative and allow it
        # (better to have a false positive than reject valid questions)
        return True

    def _validate_questions(self, questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Validate and clean extracted questions"""
        valid_questions = []

        st.info(f"🔍 Validating {len(questions)} extracted questions...")

        for i, q in enumerate(questions):
            try:
                # Check required fields
                if not isinstance(q, dict):
                    st.warning(f"Question {i+1}: Not a valid object, skipping")
                    continue

                if "question" not in q or not q["question"]:
                    st.warning(f"Question {i+1}: Missing question text, skipping")
                    continue

                if "choices" not in q or not isinstance(q["choices"], list):
                    st.warning(f"Question {i+1}: Missing or invalid choices, skipping")
                    continue

                if len(q["choices"]) < 2:
                    st.warning(f"Question {i+1}: Need at least 2 choices, skipping")
                    continue

                # Check for empty choices
                empty_choices = sum(
                    1 for c in q["choices"] if not c or (isinstance(c, str) and not c.strip())
                )
                if empty_choices > 0:
                    st.warning(f"Question {i+1}: Has {empty_choices} empty choice(s), skipping")
                    continue

                # Check for duplicate choices (exact matches)
                non_empty_choices = [
                    c for c in q["choices"] if c and (not isinstance(c, str) or c.strip())
                ]
                if len(non_empty_choices) != len(set(non_empty_choices)):
                    st.warning(f"Question {i+1}: Has duplicate choices, skipping")
                    continue

                if "correct_index" not in q:
                    st.warning(f"Question {i+1}: Missing correct_index, skipping")
                    continue

                # Check for multi-answer questions (NOT SUPPORTED)
                multi_answer_reason = self._detect_multi_answer_question(
                    q["question"],
                    q["correct_index"]
                )
                if multi_answer_reason:
                    st.warning(
                        f"Question {i+1}: Skipping - {multi_answer_reason}. "
                        f"Multi-answer questions are not supported. "
                        f"Question text: '{q['question'][:80]}...'"
                    )
                    continue

                # Validate correct_index - must be single integer
                raw_index = q["correct_index"]
                if isinstance(raw_index, list):
                    # If AI returned a list but only one answer, extract it
                    if len(raw_index) == 0:
                        st.warning(f"Question {i+1}: Empty correct_index list, skipping")
                        continue
                    elif len(raw_index) > 1:
                        # This should have been caught above, but double-check
                        st.warning(f"Question {i+1}: Multiple correct answers not supported, skipping")
                        continue
                    correct_index = int(raw_index[0])
                else:
                    correct_index = int(raw_index)

                if correct_index < 0 or correct_index >= len(q["choices"]):
                    st.warning(f"Question {i+1}: Invalid correct_index {correct_index}, skipping")
                    continue

                # Auto-validate calculations (detects and corrects AI errors)
                corrected_index, was_corrected = self._auto_validate_calculations(
                    q["question"].strip(),
                    [str(choice).strip() for choice in q["choices"]],
                    correct_index,
                )

                if was_corrected:
                    st.warning(
                        f"Question {i+1}: Auto-corrected answer from index {correct_index} to {corrected_index}"
                    )
                    correct_index = corrected_index

                # Check for OCR corruption (malformed text)
                corruption_reason = self._detect_ocr_corruption(q["question"])
                if corruption_reason:
                    st.warning(
                        f"Question {i+1}: Skipping - OCR corruption detected ({corruption_reason}). "
                        f"Question text: '{q['question'][:80]}...'"
                    )
                    continue

                # Check for incomplete questions (missing context)
                # Pass the scenario so questions with term sheets are allowed to reference "this product", etc.
                incomplete_reason = self._detect_incomplete_question(
                    q["question"],
                    scenario=q.get("scenario")
                )
                if incomplete_reason:
                    st.warning(
                        f"Question {i+1}: Skipping - {incomplete_reason}. "
                        f"Question text: '{q['question'][:80]}...'"
                    )
                    continue

                # Check if question is finance/business related (not general trivia)
                if not self._is_finance_related(q["question"], q["choices"]):
                    st.warning(
                        f"Question {i+1}: Skipping - detected as general knowledge/trivia, not finance-related. "
                        f"Question text: '{q['question'][:80]}...'"
                    )
                    continue

                # Clean and normalize question
                cleaned_question = {
                    "question": q["question"].strip(),
                    "choices": [str(choice).strip() for choice in q["choices"]],
                    "correct_index": correct_index,
                    "correct_answer": correct_index,  # Normalize to correct_answer for validation
                    "scenario": q.get("scenario", "").strip() if q.get("scenario") else None,
                    "explanation_seed": (
                        q.get("explanation_seed", "").strip() if q.get("explanation_seed") else None
                    ),
                }

                valid_questions.append(cleaned_question)

            except Exception as e:
                st.warning(f"Question {i+1}: Validation error - {str(e)}, skipping")
                continue

        # Show validation summary
        skipped = len(questions) - len(valid_questions)
        if skipped > 0:
            st.warning(f"⚠️ Validation complete: {len(valid_questions)} valid, {skipped} skipped")
        else:
            st.success(f"✅ Validation complete: All {len(valid_questions)} questions passed!")

        return valid_questions

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return ["pdf", "docx", "doc"]

    def get_file_type_help(self) -> str:
        """Get help text for supported file types"""
        return """
**Supported File Formats:**
- **PDF files (.pdf)**: Upload PDF documents containing multiple-choice questions
- **Word documents (.docx, .doc)**: Upload Word documents with questions

**Document Format Guidelines:**
The AI will automatically extract questions from your document. For best results:

1. **Clear Question Format**:
   - Number your questions (1., 2., Q1, etc.)
   - Use clear question text

2. **Answer Choices**:
   - Use consistent labeling (A, B, C, D or 1, 2, 3, 4)
   - Each choice on its own line or clearly separated

3. **Correct Answer**:
   - Indicate the correct answer (e.g., "Answer: B" or "Correct: 2")
   - Or mark with asterisk, bold, or highlight

4. **Optional Context**:
   - Include scenario/context before questions
   - Add explanations or notes (will be used for AI explanation generation)

**Example Format:**
```
Question 1: What is the capital of France?
A) London
B) Paris *
C) Berlin
D) Madrid

Answer: B
Explanation: Paris is the capital and largest city of France.
```

The AI will intelligently extract and structure all questions automatically!
"""

    def _extract_answer_key_from_images(self, doc) -> Dict[int, Dict[str, str]]:
        """
        Extract answer key from images embedded in Word document using OCR
        Returns dict mapping question number to {answer, explanation}
        """
        answer_key = {}

        try:
            # Initialize EasyOCR reader (lazy load)
            reader = easyocr.Reader(["en"], gpu=False, verbose=False)

            # Extract images from document
            images_extracted = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob
                        image = Image.open(io.BytesIO(image_data))

                        # Convert to numpy array for easyocr
                        img_array = np.array(image)

                        # Perform OCR
                        result = reader.readtext(img_array, detail=0)
                        ocr_text = " ".join(result)

                        # Parse answer sheet patterns
                        parsed_answers = self._parse_answer_sheet_text(ocr_text)
                        answer_key.update(parsed_answers)

                        images_extracted += 1

                        # Only process first few images (answer sheets usually at start)
                        if images_extracted >= 3:
                            break

                    except Exception:
                        continue

        except Exception:
            # Silently fail if OCR not available or error
            pass

        return answer_key

    def _parse_answer_sheet_text(self, text: str) -> Dict[int, Dict[str, str]]:
        """
        Parse various answer sheet formats from OCR text
        Supports patterns like:
        - "1. A  2. B  3. C"
        - "1) A  2) B  3) C"
        - "Q1: A  Q2: B"
        - Grid format with numbers and letters
        """
        answers = {}

        # Pattern 1: "1. A" or "1) A" or "1: A"
        pattern1 = r"(\d+)[\.\)\:]\s*([A-E])\b"
        matches = re.findall(pattern1, text, re.IGNORECASE)
        for q_num, answer in matches:
            answers[int(q_num)] = {"answer": answer.upper(), "explanation": ""}

        # Pattern 2: "Q1 A" or "Question 1 A"
        if not answers:
            pattern2 = r"[Qq](?:uestion)?\s*(\d+)\s+([A-E])\b"
            matches = re.findall(pattern2, text, re.IGNORECASE)
            for q_num, answer in matches:
                answers[int(q_num)] = {"answer": answer.upper(), "explanation": ""}

        # Pattern 3: Grid format - numbers on one line, letters on next
        # e.g., "1 2 3 4 5" followed by "A B C D E"
        if not answers:
            lines = text.split("\n")
            for i in range(len(lines) - 1):
                numbers = re.findall(r"\b(\d+)\b", lines[i])
                letters = re.findall(r"\b([A-E])\b", lines[i + 1], re.IGNORECASE)
                if len(numbers) == len(letters):
                    for q_num, answer in zip(numbers, letters):
                        answers[int(q_num)] = {"answer": answer.upper(), "explanation": ""}
                    if answers:
                        break

        return answers

    def _parse_word_structured(self, uploaded_file) -> List[Dict[str, Any]]:
        """
        Parse Word document using structured approach (faster and more accurate than AI)
        Looks for questions numbered like "1. Question text?" followed by choices A, B, C, D
        Also extracts answer key table if present
        """
        try:
            # Reset file pointer
            uploaded_file.seek(0)

            doc = Document(uploaded_file)
            questions = []
            answer_key = {}

            # First, try to extract answer key from images (OCR)
            if OCR_AVAILABLE:
                image_answer_key = self._extract_answer_key_from_images(doc)
                answer_key.update(image_answer_key)
                if image_answer_key:
                    st.success(
                        f"✅ OCR extracted {len(image_answer_key)} answers from answer sheet image"
                    )

            # Next, try to find answer key table (text-based)
            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    if i == 0:  # Skip header row
                        continue
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2:
                        try:
                            q_num = cells[0].strip()
                            # Extract just the number
                            q_num_match = re.match(r"(\d+)", q_num)
                            if q_num_match:
                                q_num = int(q_num_match.group(1))
                                answer = cells[1].strip()
                                explanation = cells[2].strip() if len(cells) > 2 else ""
                                answer_key[q_num] = {"answer": answer, "explanation": explanation}
                        except (ValueError, IndexError):
                            continue

            # Extract questions from paragraphs
            current_question = None
            current_choices = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Check if this is a question number (e.g., "1. Question text?")
                q_match = re.match(r"^(\d+)\.\s+(.+)", text)
                if q_match:
                    # Save previous question if exists
                    if current_question and current_choices:
                        # Convert answer to index
                        correct_index = self._convert_answer_to_index(
                            current_question.get("answer", ""), current_choices
                        )

                        questions.append(
                            {
                                "question": current_question["text"],
                                "choices": [c["text"] for c in current_choices],
                                "correct_index": correct_index,
                                "correct_answer": correct_index,  # Normalize to correct_answer for validation
                                "explanation_seed": current_question.get("explanation", ""),
                            }
                        )

                    # Start new question
                    q_num = int(q_match.group(1))
                    q_text = q_match.group(2).strip()

                    current_question = {
                        "number": q_num,
                        "text": q_text,
                        "answer": answer_key.get(q_num, {}).get("answer", ""),
                        "explanation": answer_key.get(q_num, {}).get("explanation", ""),
                    }
                    current_choices = []

                # Check if this is an answer choice (e.g., "A. Choice text")
                elif re.match(r"^[A-Z]\.\s+", text):
                    choice_match = re.match(r"^([A-Z])\.\s+(.+)", text)
                    if choice_match:
                        choice_letter = choice_match.group(1)
                        choice_text = choice_match.group(2).strip()
                        current_choices.append({"letter": choice_letter, "text": choice_text})

            # Don't forget the last question
            if current_question and current_choices:
                correct_index = self._convert_answer_to_index(
                    current_question.get("answer", ""), current_choices
                )

                questions.append(
                    {
                        "question": current_question["text"],
                        "choices": [c["text"] for c in current_choices],
                        "correct_index": correct_index,
                        "correct_answer": correct_index,  # Normalize to correct_answer for validation
                        "explanation_seed": current_question.get("explanation", ""),
                    }
                )

            return questions

        except Exception as e:
            st.warning(f"Structured parsing error: {e}")
            return []

    def _convert_answer_to_index(self, answer_str: str, choices: List[Dict[str, str]]) -> int:
        """Convert answer letters (like 'A, C') to zero-based index (takes first answer)"""
        if not answer_str:
            return 0

        # Handle multiple answers - take the first one (our system only supports single answers)
        answers = [a.strip() for a in answer_str.split(",")]
        first_answer = answers[0].upper()

        # Find the index of this choice
        for i, choice in enumerate(choices):
            if choice["letter"] == first_answer:
                return i

        return 0  # Default to first choice if not found


# Create singleton instance
document_parser = DocumentParser()

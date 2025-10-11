"""
Document ingestion and parsing system for MockExamify
Handles PDF, DOCX, CSV, and JSON file processing
"""
import json
import re
import csv
import io
import logging
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime
import docx
import PyPDF2
from models import ParsedQuestion, ParseResult, QuestionChoice

logger = logging.getLogger(__name__)


class DocumentParser:
    """Main document parser class"""
    
    def __init__(self):
        self.mcq_patterns = [
            r'^\s*(\d+)\.\s*(.+?)(?=\n\s*[A-D]\.|\n\s*\([A-D]\)|\n\s*[A-D]\)|$)',  # Question stem
            r'^\s*([A-D])\.?\s*(.+?)(?=\n\s*[A-D]\.|\n\s*\([A-D]\)|\n\s*Answer|\n\s*Correct|\n\n|$)',  # Choices
            r'^\s*(?:Answer|Correct(?:\s+Answer)?)\s*:?\s*([A-D])',  # Answer key
            r'^\s*(?:Explanation|Rationale)\s*:?\s*(.+?)(?=\n\s*\d+\.|$)'  # Explanation
        ]
    
    async def parse_file(self, file_content: bytes, file_type: str, file_name: str) -> ParseResult:
        """Parse uploaded file based on type"""
        try:
            if file_type.lower() == 'pdf':
                return await self._parse_pdf(file_content, file_name)
            elif file_type.lower() == 'docx':
                return await self._parse_docx(file_content, file_name)
            elif file_type.lower() == 'csv':
                return await self._parse_csv(file_content, file_name)
            elif file_type.lower() == 'json':
                return await self._parse_json(file_content, file_name)
            else:
                return ParseResult(
                    success=False,
                    questions=[],
                    errors=[f"Unsupported file type: {file_type}"],
                    total_parsed=0,
                    confidence_score=0.0
                )
        except Exception as e:
            logger.error(f"Error parsing file {file_name}: {e}")
            return ParseResult(
                success=False,
                questions=[],
                errors=[f"Error parsing file: {str(e)}"],
                total_parsed=0,
                confidence_score=0.0
            )
    
    async def _parse_pdf(self, file_content: bytes, file_name: str) -> ParseResult:
        """Parse PDF file for MCQ content"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return await self._extract_mcqs_from_text(text, file_name)
        except Exception as e:
            logger.error(f"Error parsing PDF {file_name}: {e}")
            return ParseResult(
                success=False,
                questions=[],
                errors=[f"PDF parsing error: {str(e)}"],
                total_parsed=0,
                confidence_score=0.0
            )
    
    async def _parse_docx(self, file_content: bytes, file_name: str) -> ParseResult:
        """Parse DOCX file for MCQ content"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return await self._extract_mcqs_from_text(text, file_name)
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_name}: {e}")
            return ParseResult(
                success=False,
                questions=[],
                errors=[f"DOCX parsing error: {str(e)}"],
                total_parsed=0,
                confidence_score=0.0
            )
    
    async def _parse_csv(self, file_content: bytes, file_name: str) -> ParseResult:
        """Parse CSV file with predefined question format"""
        try:
            content = file_content.decode('utf-8')
            csv_reader = csv.DictReader(io.StringIO(content))
            
            questions = []
            errors = []
            
            for row_num, row in enumerate(csv_reader, start=2):
                try:
                    question = self._parse_csv_row(row, row_num)
                    if question:
                        questions.append(question)
                except Exception as e:
                    errors.append(f"Row {row_num}: {str(e)}")
            
            confidence_score = 0.9 if not errors else max(0.5, 1.0 - len(errors) / len(questions)) if questions else 0.0
            
            return ParseResult(
                success=len(questions) > 0,
                questions=questions,
                errors=errors,
                total_parsed=len(questions),
                confidence_score=confidence_score
            )
        except Exception as e:
            logger.error(f"Error parsing CSV {file_name}: {e}")
            return ParseResult(
                success=False,
                questions=[],
                errors=[f"CSV parsing error: {str(e)}"],
                total_parsed=0,
                confidence_score=0.0
            )
    
    async def _parse_json(self, file_content: bytes, file_name: str) -> ParseResult:
        """Parse JSON file with predefined question format"""
        try:
            content = file_content.decode('utf-8')
            data = json.loads(content)
            
            questions = []
            errors = []
            
            # Handle different JSON structures
            if isinstance(data, list):
                question_list = data
            elif isinstance(data, dict) and 'questions' in data:
                question_list = data['questions']
            else:
                return ParseResult(
                    success=False,
                    questions=[],
                    errors=["Invalid JSON format: expected list of questions or object with 'questions' key"],
                    total_parsed=0,
                    confidence_score=0.0
                )
            
            for i, item in enumerate(question_list):
                try:
                    question = self._parse_json_question(item, i + 1)
                    if question:
                        questions.append(question)
                except Exception as e:
                    errors.append(f"Question {i + 1}: {str(e)}")
            
            confidence_score = 0.95 if not errors else max(0.7, 1.0 - len(errors) / len(questions)) if questions else 0.0
            
            return ParseResult(
                success=len(questions) > 0,
                questions=questions,
                errors=errors,
                total_parsed=len(questions),
                confidence_score=confidence_score
            )
        except json.JSONDecodeError as e:
            return ParseResult(
                success=False,
                questions=[],
                errors=[f"Invalid JSON format: {str(e)}"],
                total_parsed=0,
                confidence_score=0.0
            )
        except Exception as e:
            logger.error(f"Error parsing JSON {file_name}: {e}")
            return ParseResult(
                success=False,
                questions=[],
                errors=[f"JSON parsing error: {str(e)}"],
                total_parsed=0,
                confidence_score=0.0
            )
    
    async def _extract_mcqs_from_text(self, text: str, file_name: str) -> ParseResult:
        """Extract MCQs from plain text using pattern matching"""
        questions = []
        errors = []
        
        try:
            # Split text into potential question blocks
            question_blocks = self._split_into_question_blocks(text)
            
            for i, block in enumerate(question_blocks):
                try:
                    question = self._parse_text_block(block, i + 1)
                    if question:
                        questions.append(question)
                except Exception as e:
                    errors.append(f"Block {i + 1}: {str(e)}")
            
            # Calculate confidence based on successful parsing rate
            if not question_blocks:
                confidence_score = 0.0
            else:
                success_rate = len(questions) / len(question_blocks)
                confidence_score = min(0.8, success_rate)  # Cap at 0.8 for text parsing
            
            return ParseResult(
                success=len(questions) > 0,
                questions=questions,
                errors=errors,
                total_parsed=len(questions),
                confidence_score=confidence_score
            )
        except Exception as e:
            logger.error(f"Error extracting MCQs from text: {e}")
            return ParseResult(
                success=False,
                questions=[],
                errors=[f"Text parsing error: {str(e)}"],
                total_parsed=0,
                confidence_score=0.0
            )
    
    def _split_into_question_blocks(self, text: str) -> List[str]:
        """Split text into potential question blocks"""
        # Look for question numbers to split the text
        question_pattern = r'\n\s*(\d+)\.\s+'
        blocks = re.split(question_pattern, text)
        
        # Reconstruct blocks with question numbers
        question_blocks = []
        for i in range(1, len(blocks), 2):
            if i + 1 < len(blocks):
                question_num = blocks[i]
                question_content = blocks[i + 1]
                question_blocks.append(f"{question_num}. {question_content}")
        
        return question_blocks
    
    def _parse_text_block(self, block: str, block_num: int) -> Optional[ParsedQuestion]:
        """Parse a single text block for MCQ components"""
        try:
            lines = [line.strip() for line in block.split('\n') if line.strip()]
            
            # Extract question stem (everything before choices)
            stem_lines = []
            choice_start = -1
            
            for i, line in enumerate(lines):
                if re.match(r'^\s*[A-D]\.?\s+', line):
                    choice_start = i
                    break
                stem_lines.append(line)
            
            if not stem_lines or choice_start == -1:
                return None
            
            # Clean and join stem
            stem = ' '.join(stem_lines[1:]) if len(stem_lines) > 1 else stem_lines[0]  # Skip question number
            stem = re.sub(r'^\d+\.\s*', '', stem).strip()
            
            # Extract choices
            choices = []
            correct_index = None
            choice_pattern = r'^\s*([A-D])\.?\s+(.+?)$'
            
            for i in range(choice_start, min(choice_start + 4, len(lines))):
                match = re.match(choice_pattern, lines[i])
                if match:
                    label, text = match.groups()
                    choices.append(text.strip())
            
            if len(choices) < 4:
                return None
            
            # Look for answer key
            answer_pattern = r'(?:Answer|Correct(?:\s+Answer)?)\s*:?\s*([A-D])'
            for line in lines[choice_start + 4:]:
                match = re.search(answer_pattern, line, re.IGNORECASE)
                if match:
                    answer_letter = match.group(1).upper()
                    correct_index = ord(answer_letter) - ord('A')
                    break
            
            # Look for explanation
            explanation = None
            exp_pattern = r'(?:Explanation|Rationale)\s*:?\s*(.+)'
            for line in lines[choice_start + 4:]:
                match = re.search(exp_pattern, line, re.IGNORECASE)
                if match:
                    explanation = match.group(1).strip()
                    break
            
            # Calculate confidence
            confidence = 0.6  # Base confidence for text parsing
            if correct_index is not None:
                confidence += 0.2
            if explanation:
                confidence += 0.1
            if len(stem) > 10:  # Reasonable stem length
                confidence += 0.1
            
            return ParsedQuestion(
                stem=stem,
                choices=choices,
                correct_index=correct_index,
                explanation=explanation,
                confidence=confidence,
                raw_text=block
            )
        except Exception as e:
            logger.error(f"Error parsing text block {block_num}: {e}")
            return None
    
    def _parse_csv_row(self, row: Dict[str, str], row_num: int) -> Optional[ParsedQuestion]:
        """Parse a single CSV row into a question"""
        try:
            # Expected columns: question, choice_a, choice_b, choice_c, choice_d, correct_answer, explanation
            required_fields = ['question', 'choice_a', 'choice_b', 'choice_c', 'choice_d']
            
            for field in required_fields:
                if field not in row or not row[field].strip():
                    raise ValueError(f"Missing or empty required field: {field}")
            
            stem = row['question'].strip()
            choices = [
                row['choice_a'].strip(),
                row['choice_b'].strip(),
                row['choice_c'].strip(),
                row['choice_d'].strip()
            ]
            
            # Parse correct answer
            correct_index = None
            if 'correct_answer' in row and row['correct_answer'].strip():
                answer = row['correct_answer'].strip().upper()
                if answer in ['A', 'B', 'C', 'D']:
                    correct_index = ord(answer) - ord('A')
                elif answer in ['0', '1', '2', '3']:
                    correct_index = int(answer)
            
            explanation = row.get('explanation', '').strip() if 'explanation' in row else None
            scenario = row.get('scenario', '').strip() if 'scenario' in row else None
            
            return ParsedQuestion(
                stem=stem,
                choices=choices,
                correct_index=correct_index,
                explanation=explanation,
                scenario=scenario,
                confidence=0.9,  # High confidence for structured CSV
                raw_text=str(row)
            )
        except Exception as e:
            raise ValueError(f"Error parsing CSV row: {str(e)}")
    
    def _parse_json_question(self, item: Dict[str, Any], question_num: int) -> Optional[ParsedQuestion]:
        """Parse a single JSON question object"""
        try:
            # Required fields
            if 'question' not in item or not item['question']:
                raise ValueError("Missing 'question' field")
            
            stem = str(item['question']).strip()
            
            # Parse choices - support multiple formats
            choices = []
            if 'choices' in item and isinstance(item['choices'], list):
                choices = [str(choice).strip() for choice in item['choices']]
            elif 'options' in item and isinstance(item['options'], list):
                choices = [str(option).strip() for option in item['options']]
            else:
                # Try individual choice fields
                choice_fields = ['choice_a', 'choice_b', 'choice_c', 'choice_d']
                for field in choice_fields:
                    if field in item:
                        choices.append(str(item[field]).strip())
                
                if not choices:
                    # Try option fields
                    option_fields = ['option_a', 'option_b', 'option_c', 'option_d']
                    for field in option_fields:
                        if field in item:
                            choices.append(str(item[field]).strip())
            
            if len(choices) < 4:
                raise ValueError(f"Insufficient choices: got {len(choices)}, need 4")
            
            # Parse correct answer
            correct_index = None
            for field in ['correct_answer', 'answer', 'correct', 'correct_index']:
                if field in item and item[field] is not None:
                    answer = str(item[field]).strip().upper()
                    if answer in ['A', 'B', 'C', 'D']:
                        correct_index = ord(answer) - ord('A')
                        break
                    elif answer in ['0', '1', '2', '3']:
                        correct_index = int(answer)
                        break
                    elif field == 'correct_index' and isinstance(item[field], int):
                        correct_index = item[field]
                        break
            
            explanation = None
            for field in ['explanation', 'rationale', 'answer_explanation']:
                if field in item and item[field]:
                    explanation = str(item[field]).strip()
                    break
            
            scenario = None
            for field in ['scenario', 'context', 'case_study']:
                if field in item and item[field]:
                    scenario = str(item[field]).strip()
                    break
            
            return ParsedQuestion(
                stem=stem,
                choices=choices[:4],  # Ensure exactly 4 choices
                correct_index=correct_index,
                explanation=explanation,
                scenario=scenario,
                confidence=0.95,  # High confidence for structured JSON
                raw_text=json.dumps(item, indent=2)
            )
        except Exception as e:
            raise ValueError(f"Error parsing JSON question: {str(e)}")


class SyllabusProcessor:
    """Process syllabus documents and create searchable chunks"""
    
    def __init__(self, chunk_size: int = 1000, overlap: int = 200):
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    async def process_syllabus(self, file_content: bytes, file_type: str, 
                             file_name: str) -> Tuple[List[str], List[str]]:
        """Process syllabus document into chunks"""
        try:
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text = await self._extract_pdf_text(file_content)
            elif file_type.lower() == 'docx':
                text = await self._extract_docx_text(file_content)
            elif file_type.lower() == 'txt':
                text = file_content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type for syllabus: {file_type}")
            
            # Create chunks
            chunks = self._create_chunks(text)
            
            # Generate topic guesses for each chunk
            topic_guesses = await self._guess_topics(chunks)
            
            return chunks, topic_guesses
        except Exception as e:
            logger.error(f"Error processing syllabus {file_name}: {e}")
            return [], []
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            if chunk_text.strip():
                chunks.append(chunk_text.strip())
        
        return chunks
    
    async def _guess_topics(self, chunks: List[str]) -> List[str]:
        """Guess topics for each chunk using simple heuristics"""
        topic_guesses = []
        
        # Simple keyword-based topic guessing
        topic_keywords = {
            'risk_management': ['risk', 'volatility', 'var', 'stress test', 'portfolio risk'],
            'portfolio_management': ['portfolio', 'asset allocation', 'diversification', 'optimization'],
            'derivatives': ['options', 'futures', 'swaps', 'derivatives', 'hedging'],
            'ethics': ['ethics', 'professional conduct', 'fiduciary', 'compliance'],
            'securities': ['equity', 'bonds', 'stocks', 'securities', 'market'],
            'analysis': ['fundamental', 'technical', 'valuation', 'analysis', 'financial ratios'],
            'regulations': ['regulation', 'sfa', 'compliance', 'regulatory', 'rules'],
            'products': ['reits', 'etf', 'structured products', 'investment products']
        }
        
        for chunk in chunks:
            chunk_lower = chunk.lower()
            best_topic = 'general'
            max_score = 0
            
            for topic, keywords in topic_keywords.items():
                score = sum(1 for keyword in keywords if keyword in chunk_lower)
                if score > max_score:
                    max_score = score
                    best_topic = topic
            
            topic_guesses.append(best_topic)
        
        return topic_guesses


# Global instances
document_parser = DocumentParser()
syllabus_processor = SyllabusProcessor()


# Export functions
async def parse_document(file_content: bytes, file_type: str, file_name: str) -> ParseResult:
    """Parse document and extract questions"""
    return await document_parser.parse_file(file_content, file_type, file_name)


async def process_syllabus_document(file_content: bytes, file_type: str, 
                                  file_name: str) -> Tuple[List[str], List[str]]:
    """Process syllabus document into chunks"""
    return await syllabus_processor.process_syllabus(file_content, file_type, file_name)
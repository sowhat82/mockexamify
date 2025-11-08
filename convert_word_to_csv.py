"""
Convert Word documents with questions to CSV format for MockExamify
Handles both questions in document and answer key at the end
"""
import re
import csv
from docx import Document
import sys


def extract_questions_from_docx(docx_file):
    """Extract questions, choices, and answers from Word document"""
    doc = Document(docx_file)

    questions = []
    answer_key = {}

    # Extract all text
    full_text = '\n'.join([para.text for para in doc.paragraphs])

    # First, try to find answer key table at the end
    print("Looking for answer key table...")
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            if i == 0:  # Skip header row
                continue
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                try:
                    q_num = cells[0].strip()
                    # Extract just the number
                    q_num_match = re.match(r'(\d+)', q_num)
                    if q_num_match:
                        q_num = int(q_num_match.group(1))
                        answer = cells[1].strip()
                        explanation = cells[2].strip() if len(cells) > 2 else ""
                        answer_key[q_num] = {
                            'answer': answer,
                            'explanation': explanation
                        }
                except (ValueError, IndexError) as e:
                    continue

    print(f"Found {len(answer_key)} answers in answer key table")

    # Now extract questions
    print("\nExtracting questions...")

    # Pattern to match questions like "1. Question text?"
    question_pattern = r'(\d+)\.\s+(.+?)(?=\n[A-Z]\.|\n\d+\.|\Z)'
    choice_pattern = r'^([A-Z])\.\s+(.+?)$'

    current_question = None
    current_choices = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this is a question number
        q_match = re.match(r'^(\d+)\.\s+(.+)', text)
        if q_match:
            # Save previous question if exists
            if current_question:
                questions.append({
                    'number': current_question['number'],
                    'question': current_question['text'],
                    'choices': current_choices.copy(),
                    'answer': current_question.get('answer', ''),
                    'explanation': current_question.get('explanation', '')
                })

            # Start new question
            q_num = int(q_match.group(1))
            q_text = q_match.group(2).strip()

            current_question = {
                'number': q_num,
                'text': q_text
            }

            # Add answer from answer key if available
            if q_num in answer_key:
                current_question['answer'] = answer_key[q_num]['answer']
                current_question['explanation'] = answer_key[q_num]['explanation']

            current_choices = []

        # Check if this is an answer choice
        elif re.match(r'^[A-Z]\.\s+', text):
            choice_match = re.match(r'^([A-Z])\.\s+(.+)', text)
            if choice_match:
                choice_letter = choice_match.group(1)
                choice_text = choice_match.group(2).strip()
                current_choices.append({
                    'letter': choice_letter,
                    'text': choice_text
                })

    # Don't forget the last question
    if current_question:
        questions.append({
            'number': current_question['number'],
            'question': current_question['text'],
            'choices': current_choices.copy(),
            'answer': current_question.get('answer', ''),
            'explanation': current_question.get('explanation', '')
        })

    return questions


def convert_answer_to_index(answer_str, choices):
    """Convert answer letters (like 'A, C') to zero-based index"""
    # Handle multiple answers - for now, take the first one
    # Our system only supports single answers
    answers = [a.strip() for a in answer_str.split(',')]

    if not answers:
        return 0

    first_answer = answers[0].upper()

    # Find the index of this choice
    for i, choice in enumerate(choices):
        if choice['letter'] == first_answer:
            return i

    return 0  # Default to first choice if not found


def save_to_csv(questions, output_file):
    """Save questions to CSV format for MockExamify upload"""

    with open(output_file, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)

        # Write header
        writer.writerow([
            'question',
            'choice_1',
            'choice_2',
            'choice_3',
            'choice_4',
            'choice_5',
            'choice_6',
            'correct_index',
            'explanation'
        ])

        # Write questions
        for q in questions:
            choices = q['choices']

            # Prepare row
            row = [q['question']]

            # Add up to 6 choices (pad with empty strings if fewer)
            for i in range(6):
                if i < len(choices):
                    row.append(choices[i]['text'])
                else:
                    row.append('')

            # Add correct answer index
            correct_index = convert_answer_to_index(q['answer'], choices)
            row.append(correct_index)

            # Add explanation
            row.append(q['explanation'])

            writer.writerow(row)

    print(f"\nSaved {len(questions)} questions to {output_file}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python convert_word_to_csv.py <input.docx> [output.csv]")
        print("\nExample:")
        print("  python convert_word_to_csv.py CACS2_Mock_Paper4.docx")
        print("  python convert_word_to_csv.py CACS2_Mock_Paper4.docx paper4.csv")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '.csv')

    print(f"Converting: {input_file}")
    print(f"Output: {output_file}")
    print("="*60)

    # Extract questions
    questions = extract_questions_from_docx(input_file)

    print(f"\nExtracted {len(questions)} questions")

    # Show first question as sample
    if questions:
        q = questions[0]
        print(f"\nSample (Question 1):")
        print(f"  Question: {q['question'][:80]}...")
        print(f"  Choices: {len(q['choices'])}")
        for choice in q['choices'][:4]:
            print(f"    {choice['letter']}. {choice['text'][:60]}...")
        print(f"  Answer: {q['answer']}")
        print(f"  Explanation: {q['explanation'][:80]}...")

    # Save to CSV
    save_to_csv(questions, output_file)

    print("\n✅ Conversion complete!")
    print(f"\nNext step: Upload '{output_file}' to MockExamify")


if __name__ == "__main__":
    main()
